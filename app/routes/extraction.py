"""
Document extraction routes: pitch deck, web enrichment, financial model pipeline.
"""

import io
import json
import logging
import os
import secrets
import tempfile

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile

logger = logging.getLogger(__name__)

from ..auth import CurrentUser, get_current_user, get_optional_user
from ..database import get_db, get_model_preferences, MODEL_DEFAULTS
from ..engine.extraction import (
    extract_text_pdf_with_ocr,
    extract_text_pptx_enhanced,
    extract_tables_from_pdf,
    two_pass_extract,
)

router = APIRouter(tags=["extraction"])

# ── Shared resource list for prompts ──────────────────────────────────────────

_RESOURCES_LIST = (
    "US electricity, Global electricity, Gas to Electricity, Diesel, Gasoline, "
    "Natural Gas, Natural Gas (CCGT), Gas Turbine (CCGT), Limestone, "
    "Limestone calcination, Crushed Limestone, Li-ion Battery embodied, "
    "Li-ion Battery EV, Battery Cathode NMC62, Nickel, Polypropylene"
)


# ── Text extraction helpers ───────────────────────────────────────────────────

def _extract_text_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))
        return "\n".join(parts)[:80000]
    except ImportError:
        raise ValueError("python-docx not installed for .docx support")


def _extract_text_url(url: str) -> str:
    import requests as _req
    from bs4 import BeautifulSoup
    resp = _req.get(
        url, timeout=15,
        headers={"User-Agent": "Mozilla/5.0 (compatible; VoLo-bot/1.0)"},
    )
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)[:80000]




# ── Routes: Pitch deck / URL extraction ───────────────────────────────────────

@router.post("/api/extract")
async def extract_from_source(
    file: UploadFile = File(None),
    url: str = Form(None),
    user: CurrentUser = Depends(get_optional_user),
):
    try:
        text = ""
        source_name = ""
        tables = []

        if file and file.filename:
            data = await file.read()
            fname = file.filename.lower()
            source_name = file.filename
            if fname.endswith(".pdf"):
                text = extract_text_pdf_with_ocr(data)
                tables = extract_tables_from_pdf(data)
            elif fname.endswith((".pptx", ".ppt")):
                text = extract_text_pptx_enhanced(data)
            elif fname.endswith(".docx"):
                text = _extract_text_docx(data)
            else:
                raise HTTPException(400, "Unsupported type — upload PDF, PPTX, or DOCX")
        elif url:
            url = url.strip()
            source_name = url
            text = _extract_text_url(url)
        else:
            raise HTTPException(400, "No file or URL provided")

        if not text.strip():
            raise HTTPException(422, "No readable text found in source")

        # Resolve per-user model preference for extraction
        extraction_model = MODEL_DEFAULTS["extraction"]
        if user and user.id:
            prefs = get_model_preferences(user.id)
            extraction_model = prefs.get("extraction", extraction_model)

        result = two_pass_extract(text, source_name, tables=tables, model=extraction_model)
        return result

    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(422, str(exc))
    except Exception as exc:
        raise HTTPException(500, f"Extraction failed: {exc}")


# ── Vision-based financial extraction from screenshots ───────────────────────

_VISION_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp"}

# Canonical metric names and their common aliases (mirrors LABEL_SYNONYMS in financial_pipeline.py)
_VISION_METRIC_ALIASES: dict[str, list[str]] = {
    "revenue":        ["revenue","total revenue","net revenue","sales","net sales","gross revenue","total net revenue","total sales",
                       "service revenue","product revenue","processing revenue","tolling revenue","contract revenue","project revenue",
                       "recurring revenue"],
    "ebitda":         ["ebitda","adjusted ebitda","adj. ebitda","operating income","ebit","operating profit",
                       "adj ebitda","adjusted operating income","operating earnings"],
    "gross_profit":   ["gross profit","gross margin","gross income","gross profit margin"],
    "net_income":     ["net income","net loss","net profit","net earnings","net income/(loss)","profit/(loss)","earnings",
                       "net income (loss)","total net income","net loss attributable","loss from operations"],
    "cash":           ["cash","cash & equivalents","cash and equivalents","ending cash","cash balance","cash position",
                       "cash & cash equivalents","total cash","cash and short-term investments","available cash",
                       "cash end of period","cash eop"],
    "opex":           ["opex","operating expenses","total opex","total expenses","sg&a","r&d","operating costs",
                       "total operating expenses","selling general and administrative","g&a","general & administrative",
                       "research and development","total cost of revenue","cost of goods sold","cogs","cost of sales"],
    "gross_margin_pct": ["gross margin %","gross margin pct","gm%","gross margin percentage"],
    "capital_raised": ["capital raised","funding","total funding","capital","investment","total capital raised",
                       "total equity raised","equity raised","total invested capital"],
    "arr":            ["arr","annual recurring revenue"],
    "mrr":            ["mrr","monthly recurring revenue"],
    "customer_count": ["customers","total customers","active customers","customer count","units in field","units deployed",
                       "number of customers","active accounts","contracted customers","sites"],
    "churn_rate":     ["churn","churn rate","net churn","annual churn"],
    "runway_months":  ["runway","months runway","cash runway","months of runway"],
    "capex":          ["capex","capital expenditures","capital expenditure","property plant and equipment","ppe",
                       "investments in ppe","property and equipment"],
    "free_cash_flow": ["free cash flow","fcf","unlevered free cash flow","levered free cash flow","operating cash flow",
                       "cash from operations","net cash from operating activities"],
}

def _vision_map_metric(raw_name: str) -> str | None:
    """Map a raw label from vision output to a canonical metric key."""
    norm = raw_name.strip().lower()
    for canonical, aliases in _VISION_METRIC_ALIASES.items():
        if norm == canonical or norm in aliases:
            return canonical
        for alias in aliases:
            if alias in norm or norm in alias:
                return canonical
    return None


async def _vision_extract_financials(image_bytes: bytes, filename: str) -> dict:
    """Use Claude Vision to extract a financial model from a screenshot image.

    Returns the same dict schema as the Excel extraction route so the frontend
    and downstream pipeline are completely unaware of the difference.
    """
    import base64
    import anthropic
    import re

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        raise HTTPException(500, "ANTHROPIC_API_KEY not configured — cannot use vision extraction.")

    ext = os.path.splitext(filename)[1].lower().lstrip(".")
    media_type_map = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "webp": "image/webp"}
    media_type = media_type_map.get(ext, "image/png")
    b64_image = base64.standard_b64encode(image_bytes).decode()

    prompt = """You are a financial analyst extracting structured data from a financial model screenshot.

Carefully read every row label and column header in this image, then return ONLY a valid JSON object — no prose, no markdown fences.

JSON schema:
{
  "scale": "raw" | "thousands" | "millions" | "billions",
  "currency": "USD",
  "years": [2024, 2025, 2026],
  "metrics": {
    "<row label exactly as written>": {
      "2024": 1234567,
      "2025": null
    }
  },
  "notes": "<any important caveats, e.g. partial year, FY ending March>"
}

Rules:
1. "scale" means how the numbers are presented IN the table. Do NOT multiply — report values exactly as shown.
   I will apply the multiplier. Detect scale from headers like "($ in thousands)", "(USD M)", "in millions", etc.
2. If a year header reads "2027E", "2027F", "FY27", or "FY2027" treat it as 2027.
3. Use null for blank or unreadable cells.
4. Use negative numbers for losses (look for parentheses like "(1,234)" = -1234).
5. Include ALL rows you can read — do not filter. I will map them to canonical names.
6. Only include years between 2015 and 2040.
7. CRITICAL: rows labelled "% EBITDA Margin", "EBITDA Margin", "Gross Margin %", or any row whose values are clearly percentages (e.g. 19%, 30%) must be captured with their exact label. NEVER map percentage/margin rows onto a dollar metric like "EBITDA" or "Revenue". Keep each row label exactly as written so I can distinguish them."""

    client = anthropic.Anthropic(api_key=api_key)
    response = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=8192,
        messages=[{
            "role": "user",
            "content": [
                {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64_image}},
                {"type": "text", "text": prompt},
            ],
        }],
    )

    raw_text = response.content[0].text.strip()

    # ── Robust JSON extraction ────────────────────────────────────────────────
    # 1. Strip markdown code fences (```json ... ```)
    raw_text = re.sub(r"```[a-z]*\s*", "", raw_text)
    raw_text = raw_text.replace("```", "").strip()

    # 2. If there's still prose before/after the JSON object, extract the outermost { … }
    brace_start = raw_text.find("{")
    brace_end   = raw_text.rfind("}")
    if brace_start != -1 and brace_end != -1 and brace_end > brace_start:
        raw_text = raw_text[brace_start : brace_end + 1]

    def _repair_json(text: str) -> str:
        """Fix common LLM JSON output issues: JS comments, trailing commas, control chars."""
        # Remove JavaScript-style line comments (// ...)
        text = re.sub(r"//[^\n]*", "", text)
        # Remove JavaScript-style block comments (/* ... */)
        text = re.sub(r"/\*.*?\*/", "", text, flags=re.DOTALL)
        # Remove trailing commas before closing brace or bracket
        text = re.sub(r",(\s*[}\]])", r"\1", text)
        # Replace literal tab/newline characters inside string values that break parsing
        text = re.sub(r'(?<=")([^"]*?)(?=")', lambda m: m.group(0).replace('\n', ' ').replace('\t', ' '), text)
        return text

    try:
        extracted = json.loads(raw_text)
    except json.JSONDecodeError:
        # Attempt auto-repair before giving up
        repaired = _repair_json(raw_text)
        try:
            extracted = json.loads(repaired)
        except json.JSONDecodeError as e:
            raise HTTPException(422, f"Vision model returned unparseable JSON: {e}. Raw (first 500 chars): {raw_text[:500]}")

    # Apply scale multiplier
    scale_str = extracted.get("scale", "raw")
    scale_mult = {"raw": 1, "thousands": 1_000, "millions": 1_000_000, "billions": 1_000_000_000}.get(scale_str, 1)

    raw_metrics: dict = extracted.get("metrics", {})
    years_raw: list = extracted.get("years", [])
    valid_years = sorted({int(y) for y in years_raw if isinstance(y, (int, float)) and 2015 <= int(y) <= 2040})

    # Map raw labels → canonical metric names, apply scale
    financials: dict = {}
    unmapped: list[str] = []
    for raw_label, year_vals in raw_metrics.items():
        if not isinstance(year_vals, dict):
            continue
        canonical = _vision_map_metric(raw_label)
        if canonical is None:
            unmapped.append(raw_label)
            continue
        if canonical not in financials:
            financials[canonical] = {}
        for yr_key, val in year_vals.items():
            try:
                yr_int = int(float(str(yr_key)))
            except (ValueError, TypeError):
                continue
            if not (2015 <= yr_int <= 2040):
                continue
            if val is not None:
                try:
                    financials[canonical][str(yr_int)] = float(val) * scale_mult
                except (ValueError, TypeError):
                    pass

    # Derive fiscal_years from what was actually extracted
    all_extracted_years = sorted({int(y) for metric_vals in financials.values() for y in metric_vals.keys()})
    if not all_extracted_years:
        all_extracted_years = valid_years

    records_count = sum(len(v) for v in financials.values())

    return {
        "status": "ok",
        "file_name": filename,
        "records_count": records_count,
        "failures_count": 0,
        "financials": financials,
        "units": {},
        "fiscal_years": all_extracted_years,
        "scale_info": f"USD (scale={scale_str}, multiplier={scale_mult:,})",
        "model_summary": {
            "source": "vision_extraction",
            "manually_edited": False,
            "extraction_notes": extracted.get("notes", ""),
            "vision_model": "claude-opus-4-6",
        },
        "scenarios": None,
        "detected_scenarios": ["base"],
        "primary_scenario": "base",
        "_diagnostics": {
            "vision_extraction": True,
            "scale_detected": scale_str,
            "scale_multiplier": scale_mult,
            "unmapped_labels": unmapped,
            "raw_years": years_raw,
        },
    }


# ── Routes: Standalone financial model extraction (wizard flow) ───────────────

def _adapt_new_extractor_response(new_res: dict, file_name: str) -> dict:
    """Map the new extractor's output shape to the shape the UI expects.

    The new extractor returns:
        {scope:{sheet, years_covered, ...}, metrics:{canonical:{values, unit, ...}}, ...}

    The legacy UI expects:
        {status, file_name, financials:{metric:{year_str: value}},
         units:{metric: unit_str}, fiscal_years:[...], scale_info,
         model_summary, scenarios, detected_scenarios, primary_scenario,
         _diagnostics:{...}}
    """
    metrics = new_res.get("metrics") or {}
    scope = new_res.get("scope") or {}
    years = scope.get("years_covered") or []
    verification = new_res.get("verification") or {}

    # financials: {canonical: {year_str: value}} — drop null entries so the
    # UI's zero-filter treats missing years as missing, not zero.
    financials: dict = {}
    units: dict = {}
    for canonical, m in metrics.items():
        if not m or not isinstance(m, dict):
            continue
        vals = m.get("values") or {}
        kept = {y: v for y, v in vals.items() if v is not None}
        if not kept:
            continue
        financials[canonical] = kept
        if m.get("unit"):
            units[canonical] = m["unit"]

    # Scale info: use the unit declared on the revenue row if present.
    rev_unit = units.get("revenue") or next(iter(units.values()), None) or "USD"
    scale_info = {"USD_M": "USD_M (in millions)",
                  "USD_K": "USD_K (in thousands)",
                  "USD_B": "USD_B (in billions)",
                  "USD": "USD"}.get(rev_unit, rev_unit)

    # Surface the new extractor's selection rationale and verification
    # summary in a diagnostics block that an operator can inspect.
    diagnostics = {
        "extractor": "single_source_v0.1",
        "chosen_sheet": scope.get("sheet"),
        "scope_description": scope.get("scope_description"),
        "selection_rationale": new_res.get("selection_rationale"),
        "verification_checks_passed": verification.get("checks_passed", []),
        "verification_warnings": verification.get("warnings", []),
        "verification_errors": verification.get("errors", []),
        "candidates_considered": new_res.get("candidates_considered", []),
        "fallback_attempted": new_res.get("fallback_attempted"),
        "metric_source_cells": {
            k: m.get("source_row_excel_addr")
            for k, m in metrics.items() if m
        },
    }

    return {
        "status": new_res.get("status", "ok"),
        "file_name": file_name,
        "records_count": sum(len(v) for v in financials.values()),
        "failures_count": 0,
        "financials": financials,
        "units": units,
        "fiscal_years": years,
        "scale_info": scale_info,
        "model_summary": {
            "sheet": scope.get("sheet"),
            "description": scope.get("scope_description"),
            "years": f"{min(years)}–{max(years)}" if years else "",
            "metrics_extracted": sorted(financials.keys()),
        },
        "scenarios": None,             # new extractor does not (yet) split scenarios
        "detected_scenarios": ["base"],
        "primary_scenario": "base",
        "_diagnostics": diagnostics,
    }


@router.post("/api/extract-model")
async def extract_financial_model_standalone(
    file: UploadFile = File(...),
    user: CurrentUser = Depends(get_optional_user),
):
    """Extract financial data from an Excel/CSV model or a screenshot image."""
    ext = os.path.splitext(file.filename)[1].lower()

    # ── Image path: vision extraction ────────────────────────────────────────
    if ext in _VISION_IMAGE_EXTS:
        content = await file.read()
        return await _vision_extract_financials(content, file.filename)

    if ext not in (".xlsx", ".xlsm", ".xls", ".csv"):
        raise HTTPException(400, f"Unsupported file type: {ext}. Upload .xlsx/.xls/.csv or a screenshot (.png/.jpg/.jpeg)")

    upload_dir = os.path.join(tempfile.gettempdir(), f"fm_wiz_{secrets.token_hex(6)}")
    os.makedirs(upload_dir, exist_ok=True)
    input_path = os.path.join(upload_dir, file.filename)

    content = await file.read()
    with open(input_path, "wb") as f:
        f.write(content)

    out_dir = os.path.join(upload_dir, "output")
    os.makedirs(out_dir, exist_ok=True)

    # ── Try new single-source-of-truth extractor first (opt-out via env) ─────
    # The new extractor guarantees every row comes from one sheet, so there's
    # no silent source-switch bug. If it returns ok/ok_with_warnings we use
    # its output. Any error or ambiguity falls back to the legacy pipeline so
    # behavior on non-standard workbooks is preserved.
    if os.environ.get("VOLO_USE_NEW_EXTRACTOR", "1") != "0":
        try:
            from ..engine.extract_financials import extract as _new_extract
            new_res = _new_extract(input_path)
            if new_res.get("status") in ("ok", "ok_with_warnings"):
                return _adapt_new_extractor_response(
                    new_res, file_name=file.filename,
                )
            # clarifying_question or error -> fall through to legacy pipeline
            logger.info(
                "New extractor did not resolve (status=%s); falling back to legacy",
                new_res.get("status"),
            )
        except Exception as _new_err:
            logger.exception("New extractor raised; falling back to legacy: %s", _new_err)

    try:
        from ..engine.financial_pipeline import run_pipeline
        result = run_pipeline(
            input_path=input_path,
            company_id="wizard",
            fy_end_month=12,
            out_dir=out_dir,
        )

        raw_years = result.get("fiscal_years", [])
        valid_years = sorted([y for y in raw_years if isinstance(y, (int, float)) and 2015 <= y <= 2050])

        def _year_ok(k):
            try:
                return 2015 <= int(float(k)) <= 2050
            except (ValueError, TypeError):
                return False

        records = result.get("records", [])
        failures = result.get("failures", [])
        from ..engine.financial_pipeline import UNIT_METRICS

        scale_info = "USD"

        # Keywords that indicate a consolidated / company-level sheet (highest priority)
        CONSOLIDATED_KEYWORDS = {"consolidated", "consol", "total", "combined", "group"}
        # Keywords that indicate a forecast / summary sheet (high priority)
        FORECAST_KEYWORDS = {"forecast", "outputs", "summary",
                             "projections", "pro forma", "proforma", "model",
                             "p&l", "income statement", "income_statement", "financials"}
        # Keywords that indicate a segment / geography / entity sheet (low priority)
        SEGMENT_KEYWORDS = {
            "north america", "europe", "asia", "apac", "emea", "latam",
            "norway", "usa", "uk", "germany", "japan", "china", "india",
            "canada", "australia", "france", "brazil", "mexico",
            "segment", "division", "region", "subsidiary", "entity",
        }

        def _sheet_priority(sheet_name):
            """Return priority score: higher = preferred. Consolidated > Forecast > Default > Segment."""
            sl = (sheet_name or "").lower()
            is_consolidated = any(kw in sl for kw in CONSOLIDATED_KEYWORDS)
            is_forecast = any(kw in sl for kw in FORECAST_KEYWORDS)
            is_segment = any(kw in sl for kw in SEGMENT_KEYWORDS)

            if is_consolidated:
                return 100 + (10 if is_forecast else 0)  # "Consolidated Forecast" = 110
            if is_segment:
                return 10  # Always low priority regardless of other keywords
            if is_forecast:
                return 80
            return 50  # Default/unknown sheets

        best_records = {}
        for r in records:
            fy = r.get("fiscal_year")
            if not _year_ok(fy):
                continue
            metric = r.get("metric")
            val = r.get("value_usd")
            if val is None:
                continue

            scenario = r.get("scenario", "base")

            prov = r.get("provenance", {})
            sheet = prov.get("source_sheet", "")

            key = (scenario, metric, fy)
            priority = _sheet_priority(sheet)
            existing = best_records.get(key)
            if existing is None:
                best_records[key] = {"val": val, "sheet": sheet, "priority": priority, "record": r}
            elif priority > existing["priority"]:
                best_records[key] = {"val": val, "sheet": sheet, "priority": priority, "record": r}

        # Build per-scenario financials and units
        all_scenarios_financials = {}
        all_scenarios_units = {}
        detected_scenarios = set()

        for (scenario, metric, fy), info in best_records.items():
            detected_scenarios.add(scenario)
            fy_str = str(fy)
            r = info["record"]
            if metric in UNIT_METRICS:
                all_scenarios_units.setdefault(scenario, {}).setdefault(metric, {})[fy_str] = {
                    "value": info["val"],
                    "unit_type": r.get("unit_type"),
                }
            else:
                all_scenarios_financials.setdefault(scenario, {}).setdefault(metric, {})[fy_str] = info["val"]

        def _clean_financials(fin_dict):
            """Remove margin_pct_annual, filter outliers relative to revenue."""
            rev_by_year = {}
            if "revenue" in fin_dict:
                rev_by_year = {k: abs(v) for k, v in fin_dict["revenue"].items() if v}
            rev_max = max(rev_by_year.values()) if rev_by_year else 0

            for metric in list(fin_dict.keys()):
                if metric == "revenue":
                    continue
                if metric == "margin_pct_annual":
                    del fin_dict[metric]
                    continue
                series = fin_dict[metric]
                vals = [v for v in series.values() if v is not None and v != 0]
                if not vals:
                    del fin_dict[metric]
                    continue
                cleaned = {}
                for k, v in series.items():
                    if v is None:
                        continue
                    absv = abs(v)
                    same_yr_rev = rev_by_year.get(k, 0)
                    if same_yr_rev > 0 and absv > same_yr_rev * 50:
                        continue
                    if same_yr_rev == 0 and rev_max > 0 and absv > rev_max * 50:
                        continue
                    if rev_max > 0 and absv > rev_max * 50:
                        continue
                    cleaned[k] = v
                fin_dict[metric] = cleaned
                if not cleaned:
                    del fin_dict[metric]
            return fin_dict

        # Clean each scenario
        for sc in list(all_scenarios_financials.keys()):
            all_scenarios_financials[sc] = _clean_financials(all_scenarios_financials[sc])

        # Primary = "base" scenario; fallback to scenario with most data
        primary_scenario = "base"
        if "base" not in all_scenarios_financials and all_scenarios_financials:
            primary_scenario = max(
                all_scenarios_financials,
                key=lambda s: sum(len(v) for v in all_scenarios_financials[s].values())
            )
        financials = all_scenarios_financials.get(primary_scenario, {})
        units = all_scenarios_units.get(primary_scenario, {})

        # Build scenarios output (bear/base/bull)
        scenarios_output = {}
        for sc in sorted(detected_scenarios):
            scenarios_output[sc] = {
                "financials": all_scenarios_financials.get(sc, {}),
                "units": all_scenarios_units.get(sc, {}),
            }

        summary = result.get("model_summary", {})
        clean_summary = {}
        for k, v in summary.items():
            if isinstance(v, dict):
                clean_summary[k] = json.dumps(v)
            elif isinstance(v, list):
                clean_summary[k] = ", ".join(str(x) for x in v)
            else:
                clean_summary[k] = v

        # Build extraction diagnostics: which sheets had which metrics, and priority scores
        _diag_sheets = {}
        for r in records:
            prov = r.get("provenance", {})
            sheet = prov.get("source_sheet", "?")
            metric = r.get("metric", "?")
            fy = r.get("fiscal_year")
            if sheet not in _diag_sheets:
                _diag_sheets[sheet] = {"priority": _sheet_priority(sheet), "metrics": {}}
            _diag_sheets[sheet]["metrics"].setdefault(metric, []).append(fy)

        _diag_best = {}
        for (sc, metric, fy), info in best_records.items():
            _diag_best.setdefault(metric, []).append({
                "fy": fy, "sheet": info["sheet"], "priority": info["priority"],
                "val": round(info["val"], 2) if info["val"] else None
            })

        return {
            "status": "ok",
            "file_name": file.filename,
            "records_count": len([r for r in records if _year_ok(r.get("fiscal_year", 0))]),
            "failures_count": result.get("failures_count", 0),
            "financials": financials,
            "units": units,
            "fiscal_years": valid_years,
            "scale_info": scale_info,
            "model_summary": clean_summary,
            "scenarios": scenarios_output if len(detected_scenarios) > 1 else None,
            "detected_scenarios": sorted(detected_scenarios),
            "primary_scenario": primary_scenario,
            "_diagnostics": {
                "sheets_processed": _diag_sheets,
                "best_record_sources": _diag_best,
                "pipeline_sheet_diagnostics": result.get("_sheet_diagnostics", []),
                "all_workbook_sheets": result.get("_all_sheets", []),
                "extraction_failures": [
                    {
                        "metric": f.get("metric"),
                        "sheet": f.get("provenance", {}).get("source_sheet"),
                        "fail_code": f.get("quality", {}).get("fail_code"),
                        "details": f.get("quality", {}).get("details", "")[:200],
                        "has_formulas": f.get("quality", {}).get("_has_formulas", False),
                    }
                    for f in failures
                ],
            },
        }
    except Exception as exc:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Financial model extraction failed: {exc}")


